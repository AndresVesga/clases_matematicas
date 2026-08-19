# -*- coding: utf-8 -*-
"""
Convierte las clases en Markdown a documentos Word (.docx) editables.

Uso (desde la carpeta raiz del proyecto):

    python scripts/md_a_word.py

Eso genera:

    entregables/word/clase_01.docx ... clase_06.docx
    entregables/word/Matematicas_Unidad_Numeros_Reales_y_Complejos.docx

Opciones:

    python scripts/md_a_word.py clases/clase_03.md      -> convierte solo ese archivo
    python scripts/md_a_word.py --sin-consolidado       -> no genera el documento unico

Requisito unico: la libreria python-docx.
    pip install python-docx

El script entiende el subconjunto de Markdown que usan las clases:
titulos (#, ##, ###), parrafos, negrita (**texto**), listas con viñetas (*),
listas numeradas (1. 2. 3.) y tablas simples con |.
No usa lineas divisorias porque el proyecto no las utiliza.

MATEMATICAS
Lo que va entre $` y `$ se convierte en una ecuacion de verdad de Word,
con fracciones, raices y exponentes bien formados. Por ejemplo,
$`\\frac{2}{3}`$ sale como una fraccion con su raya horizontal.
Se usa ese delimitador, y no el signo de pesos, por dos razones: GitHub lo
renderiza igual, y no choca con los precios en pesos de los enunciados.
"""

import io
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from latex_a_omml import omml_de

RAIZ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CARPETA_CLASES = os.path.join(RAIZ, "clases")
CARPETA_SALIDA = os.path.join(RAIZ, "entregables", "word")
NOMBRE_CONSOLIDADO = "Matematicas_Unidad_Numeros_Reales_y_Complejos.docx"
TITULO_CONSOLIDADO = "Matematicas. Numeros reales y numeros complejos"

RE_TABLA_SEP = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)*\|?\s*$")
RE_VINETA = re.compile(r"^(\s*)[-*]\s+(.*)$")
RE_NUMERO = re.compile(r"^\s*(\d+)\.\s+(.*)$")
# Formula en medio de una frase: $`...`$
# Es el mismo delimitador que GitHub sabe renderizar, y no choca con los
# precios en pesos, que aparecen en casi todos los enunciados.
RE_MATE = re.compile(r"\$`(.+?)`\$", re.S)

# Tipografia del documento. Cambria es la serif que Word trae de fabrica y
# es la pareja natural de Cambria Math, la fuente de las ecuaciones: asi el
# texto y las formulas se ven como un mismo documento y no como un pegote.
FUENTE = "Cambria"


def configurar_estilos(doc):
    """Tipografia comoda de leer e imprimir."""
    normal = doc.styles["Normal"]
    normal.font.name = FUENTE
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for nombre, tam in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        estilo = doc.styles[nombre]
        estilo.font.name = FUENTE
        estilo.font.size = Pt(tam)
        estilo.font.bold = True
        estilo.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)
        estilo.paragraph_format.space_before = Pt(12)
        estilo.paragraph_format.space_after = Pt(4)
        estilo.paragraph_format.keep_with_next = True


def escribir_plano(parrafo, texto):
    """Escribe texto normal respetando la negrita **asi**."""
    for i, trozo in enumerate(texto.split("**")):
        if not trozo:
            continue
        run = parrafo.add_run(trozo)
        run.bold = (i % 2 == 1)


def escribir_texto(parrafo, texto):
    """Escribe una linea mezclando texto normal y ecuaciones.

    Todo lo que este entre $` y `$ se inserta como ecuacion de Word.
    """
    posicion = 0
    for formula in RE_MATE.finditer(texto):
        escribir_plano(parrafo, texto[posicion:formula.start()])
        parrafo._p.append(omml_de(formula.group(1)))
        posicion = formula.end()
    escribir_plano(parrafo, texto[posicion:])


def agregar_tabla(doc, filas):
    """filas es una lista de listas de celdas ya limpias."""
    tabla = doc.add_table(rows=len(filas), cols=len(filas[0]))
    tabla.style = "Table Grid"
    for f, fila in enumerate(filas):
        for c, celda in enumerate(fila):
            if c >= len(tabla.columns):
                continue
            parrafo = tabla.cell(f, c).paragraphs[0]
            escribir_texto(parrafo, celda)
            if f == 0:
                for run in parrafo.runs:
                    run.bold = True
    doc.add_paragraph()


def celdas(linea):
    return [c.strip() for c in linea.strip().strip("|").split("|")]


def volcar_markdown(doc, texto, desplazar_titulos=0):
    """Escribe el contenido del markdown dentro del documento.

    desplazar_titulos=1 baja un nivel cada titulo; se usa en el consolidado
    para que el titulo del documento quede por encima del de cada clase.
    """
    lineas = texto.split("\n")
    i = 0
    while i < len(lineas):
        linea = lineas[i].rstrip()

        if not linea.strip():
            i += 1
            continue

        # Tabla: la linea actual tiene | y la siguiente es el separador ---|---
        if "|" in linea and i + 1 < len(lineas) and RE_TABLA_SEP.match(lineas[i + 1]):
            filas = [celdas(linea)]
            i += 2
            while i < len(lineas) and "|" in lineas[i] and lineas[i].strip():
                filas.append(celdas(lineas[i]))
                i += 1
            agregar_tabla(doc, filas)
            continue

        # Titulos
        if linea.startswith("#"):
            nivel = len(linea) - len(linea.lstrip("#"))
            contenido = linea.lstrip("#").strip()
            nivel = min(nivel + desplazar_titulos, 4)
            parrafo = doc.add_heading("", level=nivel)
            escribir_texto(parrafo, contenido)
            i += 1
            continue

        # Vinetas. Una viñeta indentada (los items a, b, c de un ejercicio)
        # queda como sublista, un nivel adentro.
        m = RE_VINETA.match(linea)
        if m:
            estilo = "List Bullet 2" if len(m.group(1)) >= 2 else "List Bullet"
            parrafo = doc.add_paragraph(style=estilo)
            escribir_texto(parrafo, m.group(2))
            i += 1
            continue

        # Listas numeradas: se conserva el numero tal cual esta escrito,
        # para que la actividad siga yendo del 1 al 15 sin reiniciarse.
        m = RE_NUMERO.match(linea)
        if m:
            parrafo = doc.add_paragraph()
            parrafo.paragraph_format.left_indent = Pt(24)
            parrafo.paragraph_format.first_line_indent = Pt(-12)
            escribir_texto(parrafo, m.group(1) + ". " + m.group(2))
            i += 1
            continue

        # Parrafo normal
        parrafo = doc.add_paragraph()
        escribir_texto(parrafo, linea.strip())
        i += 1


def leer(ruta):
    with io.open(ruta, encoding="utf-8") as f:
        return f.read()


def convertir(ruta_md, ruta_docx):
    doc = Document()
    configurar_estilos(doc)
    volcar_markdown(doc, leer(ruta_md))
    doc.save(ruta_docx)
    print("Generado:", os.path.relpath(ruta_docx, RAIZ))


def consolidar(rutas_md, ruta_docx):
    doc = Document()
    configurar_estilos(doc)
    portada = doc.add_heading("", level=0)
    escribir_texto(portada, TITULO_CONSOLIDADO)
    portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    escribir_texto(sub, "Seis clases de 80 minutos. Material para el docente.")
    for ruta in rutas_md:
        doc.add_page_break()
        volcar_markdown(doc, leer(ruta), desplazar_titulos=0)
    doc.save(ruta_docx)
    print("Generado:", os.path.relpath(ruta_docx, RAIZ))


def main():
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    hacer_consolidado = "--sin-consolidado" not in sys.argv

    if args:
        rutas = [os.path.abspath(a) for a in args]
        hacer_consolidado = False
    else:
        rutas = sorted(
            os.path.join(CARPETA_CLASES, n)
            for n in os.listdir(CARPETA_CLASES)
            if n.endswith(".md")
        )

    if not os.path.isdir(CARPETA_SALIDA):
        os.makedirs(CARPETA_SALIDA)

    for ruta in rutas:
        nombre = os.path.splitext(os.path.basename(ruta))[0] + ".docx"
        convertir(ruta, os.path.join(CARPETA_SALIDA, nombre))

    if hacer_consolidado:
        consolidar(rutas, os.path.join(CARPETA_SALIDA, NOMBRE_CONSOLIDADO))

    print("Listo. Los documentos estan en entregables/word/")


if __name__ == "__main__":
    main()
